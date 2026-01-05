#!/usr/bin/env python3
"""
将原有技术方案和补充章节合并，并转换为完整的Word文档
"""
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE

def read_markdown_file(file_path):
    """读取Markdown文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def merge_technical_docs(original_md, supplement_md):
    """
    合并原有技术方案和补充章节
    
    原有技术方案结构：
    1. 项目概述
    2. 总体架构设计
    3. 技术栈选型
    4. 核心功能模块设计
    5. 数据库设计概览（不完整）
    6. API设计规范
    7. 部署与运维（简略）
    8. 总结
    
    补充章节结构：
    五、数据库设计（完整）
    七、界面设计与实现（完整）
    九、部署与发布（完整）
    十一、智能体设计专章（完整）
    
    合并策略：
    1. 用补充章节的第五章替换原有第五章
    2. 在原有第六章后插入补充章节的第七章
    3. 在原有第七章后插入补充章节的第九章
    4. 在原有第八章前插入补充章节的第十一章
    """
    
    # 分割原有技术方案为章节
    original_lines = original_md.split('\n')
    
    # 查找各章节位置
    chapter_positions = {}
    for i, line in enumerate(original_lines):
        if re.match(r'^## \d+\.', line) or re.match(r'^# ', line):
            chapter_title = line.strip()
            chapter_positions[chapter_title] = i
    
    # 提取补充章节内容
    supplement_chapters = {}
    current_chapter = None
    chapter_content = []
    
    for line in supplement_md.split('\n'):
        # 检测章节标题（如 "## 五、数据库设计"）
        chapter_match = re.match(r'^## (.*)$', line)
        if chapter_match:
            if current_chapter and chapter_content:
                supplement_chapters[current_chapter] = '\n'.join(chapter_content)
                chapter_content = []
            current_chapter = chapter_match.group(1)
        elif current_chapter:
            chapter_content.append(line)
    
    # 添加最后一个章节
    if current_chapter and chapter_content:
        supplement_chapters[current_chapter] = '\n'.join(chapter_content)
    
    # 构建合并后的内容
    merged_lines = []
    i = 0
    while i < len(original_lines):
        line = original_lines[i]
        
        # 检查是否到了需要替换或插入的章节
        if line == '## 5. 数据库设计概览':
            # 替换为完整的第五章
            merged_lines.append('## 5. 数据库设计')
            merged_lines.append(supplement_chapters.get('五、数据库设计', ''))
            # 跳过原有第五章内容
            while i < len(original_lines) and not re.match(r'^## 6\.', original_lines[i]):
                i += 1
            continue
        
        elif line == '## 6. API 设计规范':
            # 添加原有第六章
            merged_lines.append(line)
            j = i + 1
            while j < len(original_lines) and not re.match(r'^## 7\.', original_lines[j]):
                merged_lines.append(original_lines[j])
                j += 1
            i = j
            
            # 插入补充章节第七章
            merged_lines.append('')
            merged_lines.append('## 7. 界面设计与实现')
            merged_lines.append(supplement_chapters.get('七、界面设计与实现', ''))
            continue
        
        elif line == '## 7. 部署与运维':
            # 添加原有第七章
            merged_lines.append(line)
            j = i + 1
            while j < len(original_lines) and not re.match(r'^## 8\.', original_lines[j]):
                merged_lines.append(original_lines[j])
                j += 1
            i = j
            
            # 插入补充章节第九章
            merged_lines.append('')
            merged_lines.append('## 9. 部署与发布')
            merged_lines.append(supplement_chapters.get('九、部署与发布', ''))
            
            # 插入补充章节第十一章
            merged_lines.append('')
            merged_lines.append('## 11. 智能体设计专章')
            merged_lines.append(supplement_chapters.get('十一、智能体设计专章', ''))
            continue
        
        else:
            merged_lines.append(line)
            i += 1
    
    return '\n'.join(merged_lines)

def convert_markdown_to_docx(md_content, docx_file_path):
    """将Markdown内容转换为Word文档"""
    lines = md_content.split('\n')
    doc = Document()
    
    # 添加自定义样式
    styles = doc.styles
    
    # 代码样式
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.right_indent = Inches(0.5)
    
    # 处理状态变量
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        line = line.rstrip('\n')
        
        # 处理代码块
        if line.strip().startswith('```'):
            if in_code_block:
                code_para = doc.add_paragraph(style='Code')
                code_para.add_run('\n'.join(code_block_content))
                in_code_block = False
                code_block_content = []
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # 处理标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=min(level, 6))
            continue
        
        # 处理列表
        if re.match(r'^[\s]*[-*+]\s+', line):
            text = re.sub(r'^[\s]*[-*+]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
            continue
        
        if re.match(r'^[\s]*\d+\.\s+', line):
            text = re.sub(r'^[\s]*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
            continue
        
        # 处理表格（简化处理，仅识别简单表格）
        if '|' in line and line.strip().startswith('|'):
            # 简单表格处理
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            if not hasattr(doc, '_current_table'):
                doc._current_table = doc.add_table(rows=0, cols=len(cells))
                doc._current_table.style = 'Table Grid'
            
            row = doc._current_table.add_row()
            for j, cell_content in enumerate(cells):
                if j < len(row.cells):
                    row.cells[j].text = cell_content
            continue
        elif hasattr(doc, '_current_table'):
            # 表格结束
            del doc._current_table
            doc.add_paragraph()
        
        # 空行
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # 普通段落
        doc.add_paragraph(line)
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"已成功生成完整技术方案: {docx_file_path}")

def main():
    # 读取文件
    original_md = read_markdown_file('技术方案.md')
    supplement_md = read_markdown_file('技术方案_补充章节.md')
    
    # 合并文档
    merged_md = merge_technical_docs(original_md, supplement_md)
    
    # 保存合并后的Markdown文件（可选）
    with open('完整技术方案.md', 'w', encoding='utf-8') as f:
        f.write(merged_md)
    
    # 转换为Word文档
    convert_markdown_to_docx(merged_md, '完整技术方案.docx')
    
    print("文档合并与转换完成！")
    print("生成的文件:")
    print("1. 完整技术方案.md - 合并后的Markdown文档")
    print("2. 完整技术方案.docx - 完整的Word格式技术方案")

if __name__ == '__main__':
    main()
