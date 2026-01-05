#!/usr/bin/env python3
"""
检查Word文档中数据库字段名的下划线是否保留
"""
from docx import Document

def check_field_names(docx_file):
    """检查Word文档中的字段名"""
    doc = Document(docx_file)
    
    # 要检查的关键字段名
    target_fields = [
        'ua_id', 'ua_username', 'ua_email', 'ua_phone',
        'ua_password_hash', 'ua_status', 'ua_invite_code',
        'ls_id', 'ls_user_id', 'ls_name', 'ls_description',
        'rs_id', 'rs_learning_space_id', 'rs_title',
        'rc_id', 'rc_resource_id', 'rc_content'
    ]
    
    print(f"检查文件: {docx_file}")
    print("=" * 60)
    
    found_fields = []
    missing_fields = []
    
    # 检查所有段落
    for para in doc.paragraphs:
        text = para.text
        if text:
            for field in target_fields:
                if field in text and field not in found_fields:
                    found_fields.append(field)
    
    # 检查表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if text:
                    for field in target_fields:
                        if field in text and field not in found_fields:
                            found_fields.append(field)
    
    # 找出缺失的字段
    for field in target_fields:
        if field not in found_fields:
            missing_fields.append(field)
    
    print(f"找到的字段 ({len(found_fields)}个):")
    for field in sorted(found_fields):
        print(f"  ✓ {field}")
    
    print(f"\n未找到的字段 ({len(missing_fields)}个):")
    for field in sorted(missing_fields):
        print(f"  ✗ {field}")
    
    print("\n" + "=" * 60)
    print("检查完成!")
    
    # 显示一些示例文本
    print("\n示例内容（前5段）:")
    for i, para in enumerate(doc.paragraphs[:5]):
        if para.text.strip():
            print(f"[段落 {i+1}]: {para.text[:80]}...")
    
    return len(found_fields) > 0

if __name__ == '__main__':
    files_to_check = [
        '技术方案_补充章节_clean_fixed.docx',
        '技术方案_补充章节.docx'
    ]
    
    for file in files_to_check:
        import os
        if os.path.exists(file):
            print("\n" + "=" * 60)
            check_field_names(file)
        else:
            print(f"\n文件不存在: {file}")
