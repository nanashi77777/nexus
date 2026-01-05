#!/usr/bin/env python3
"""
将Markdown文件转换为干净的Word文档，去除所有Markdown格式标记
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_markdown_format(text):
    """
    清理Markdown格式标记，返回纯文本
    
    处理规则：
    1. **粗体** → 粗体
    2. *斜体* → 斜体  
    3. `代码` → 代码
    4. [链接](URL) → 链接
    5. ![图片](URL) → [图片]
    6. # 标题 → 标题（去除#号）
    7. ~~删除线~~ → 删除线
    8. > 引用 → 引用（去除>号）
    """
    if not text:
        return text
    
    # 保存原始文本用于特殊处理
    original = text
    
    # 1. 处理行内代码 `code`
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # 2. 处理粗体 **bold** 或 __bold__
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    
    # 3. 处理斜体 *italic*（完全保留单下划线，因为数据库字段名使用下划线）
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # 注意：不移除单下划线，以免影响数据库字段名（如 ua_id, ls_user_id 等）
    # text = re.sub(r'_([^_\s]+)_', r'\1', text)
    
    # 4. 处理删除线 ~~strikethrough~~
    text = re.sub(r'~~([^~]+)~~', r'\1', text)
    
    # 5. 处理链接 [text](url)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    
    # 6. 处理图片 ![alt](url)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'[图片]', text)
    
    # 7. 处理引用标记 > 
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        # 去除开头的 > 和空格
        line = re.sub(r'^>\s*', '', line)
        cleaned_lines.append(line)
    text = '\n'.join(cleaned_lines)
    
    # 8. 处理标题标记 # ## ###
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    
    # 9. 处理无序列表标记 - * +
    text = re.sub(r'^[\s]*[-*+]\s+', '', text, flags=re.MULTILINE)
    
    # 10. 处理有序列表标记 1. 2. 3.
    text = re.sub(r'^[\s]*\d+\.\s+', '', text, flags=re.MULTILINE)
    
    # 11. 处理代码块标记 ```language 和 ```
    text = re.sub(r'^```.*$', '', text, flags=re.MULTILINE)
    
    # 12. 处理表格分隔符 |---|
    text = re.sub(r'^\|[-:| ]+\|$', '', text, flags=re.MULTILINE)
    
    return text

def convert_markdown_to_clean_docx(md_file_path, docx_file_path):
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 创建Word文档
    doc = Document()
    
    # 添加自定义样式
    styles = doc.styles
    
    # 代码样式
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)
        code_style.font.color.rgb = RGBColor(0x36, 0x41, 0x46)  # 深灰色
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.right_indent = Inches(0.5)
        code_style.paragraph_format.line_spacing = 1.15
    
    # 表格样式
    if 'Table' not in styles:
        table_style = styles.add_style('Table', WD_STYLE_TYPE.PARAGRAPH)
        table_style.font.name = 'Calibri'
        table_style.font.size = Pt(9)
    
    # 处理状态变量
    in_code_block = False
    in_table = False
    code_block_content = []
    table_rows = []
    current_code_language = ""
    
    # 处理每一行
    for i, line in enumerate(lines):
        line = line.rstrip('\n')
        original_line = line
        
        # 处理代码块开始
        if line.strip().startswith('```'):
            if not in_code_block:
                # 开始代码块
                in_code_block = True
                code_block_content = []
                language_match = re.match(r'^```(\w+)', line.strip())
                current_code_language = language_match.group(1) if language_match else ""
            else:
                # 结束代码块
                if code_block_content:
                    # 添加代码块标题（如果有语言）
                    if current_code_language:
                        code_title = doc.add_paragraph(f"代码示例 ({current_code_language}):")
                        code_title.runs[0].bold = True
                    
                    # 添加代码内容
                    code_para = doc.add_paragraph(style='Code')
                    code_text = '\n'.join(code_block_content)
                    code_run = code_para.add_run(code_text)
                    code_run.font.name = 'Consolas'
                    
                    # 添加空行
                    doc.add_paragraph()
                
                in_code_block = False
                code_block_content = []
                current_code_language = ""
            continue
        
        # 如果在代码块中，收集内容（不清理格式）
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # 处理表格行
        if '|' in line and (line.strip().startswith('|') or '|' in line[1:]):
            if not in_table and len(line.split('|')) > 2:  # 至少有两列
                in_table = True
                table_rows = []
            
            if in_table:
                # 清理表格行中的Markdown格式
                cleaned_cells = []
                cells = [cell.strip() for cell in line.strip('|').split('|')]
                for cell in cells:
                    cleaned_cell = clean_markdown_format(cell)
                    cleaned_cells.append(cleaned_cell)
                table_rows.append(cleaned_cells)
            continue
        
        # 表格结束，处理表格
        elif in_table and table_rows:
            _process_clean_table(doc, table_rows)
            in_table = False
            table_rows = []
        
        # 处理标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_markdown_format(heading_match.group(2))
            doc.add_heading(text, level=min(level, 6))
            continue
        
        # 处理列表项（清理后）
        if re.match(r'^[\s]*[-*+]\s+', line):
            # 无序列表
            text = re.sub(r'^[\s]*[-*+]\s+', '', line)
            cleaned_text = clean_markdown_format(text)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(cleaned_text)
            continue
        
        if re.match(r'^[\s]*\d+\.\s+', line):
            # 有序列表
            text = re.sub(r'^[\s]*\d+\.\s+', '', line)
            cleaned_text = clean_markdown_format(text)
            p = doc.add_paragraph(style='List Number')
            p.add_run(cleaned_text)
            continue
        
        # 处理引用块
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line)
            cleaned_text = clean_markdown_format(text)
            p = doc.add_paragraph()
            p.add_run(cleaned_text).italic = True
            p.paragraph_format.left_indent = Inches(0.5)
            continue
        
        # 处理空行
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # 普通段落 - 清理所有Markdown格式
        cleaned_text = clean_markdown_format(line)
        if cleaned_text.strip():  # 只添加非空内容
            # 检查是否是连续文本的一部分
            if i > 0 and lines[i-1].strip() and not lines[i-1].strip().startswith(('#', '-', '*', '+', '>', '|', '`')):
                # 续接上一段
                if doc.paragraphs:
                    last_para = doc.paragraphs[-1]
                    if last_para.text:
                        last_para.add_run(" " + cleaned_text)
                    else:
                        last_para.add_run(cleaned_text)
                else:
                    doc.add_paragraph(cleaned_text)
            else:
                doc.add_paragraph(cleaned_text)
    
    # 处理最后可能存在的表格
    if in_table and table_rows:
        _process_clean_table(doc, table_rows)
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"已成功将 {md_file_path} 转换为干净的 {docx_file_path}")
    print("所有Markdown格式标记已清理完毕")

def _process_clean_table(doc, table_rows):
    """处理清理后的表格"""
    if not table_rows or len(table_rows) < 2:
        return
    
    # 过滤空行
    valid_rows = []
    for row in table_rows:
        if any(cell.strip() for cell in row):
            valid_rows.append(row)
    
    if len(valid_rows) < 2:
        return
    
    # 创建Word表格
    table = doc.add_table(rows=len(valid_rows), cols=len(valid_rows[0]))
    table.style = 'Table Grid'
    
    # 填充表格内容
    for i, row_cells in enumerate(valid_rows):
        for j, cell_content in enumerate(row_cells):
            if j < len(table.rows[i].cells):
                table.cell(i, j).text = cell_content
    
    # 添加空行分隔
    doc.add_paragraph()

def process_supplement_chapter():
    """专门处理补充章节文件"""
    input_file = '技术方案_补充章节.md'
    output_file = '技术方案_补充章节_clean_fixed.docx'
    
    print(f"正在处理补充章节: {input_file}")
    print("清理Markdown格式标记中...")
    
    try:
        convert_markdown_to_clean_docx(input_file, output_file)
        print(f"✓ 清理完成: {output_file}")
        
        # 验证文件
        import os
        if os.path.exists(output_file):
            file_size = os.path.getsize(output_file) / 1024
            print(f"✓ 文件大小: {file_size:.1f} KB")
            return True
        else:
            print("✗ 文件创建失败")
            return False
            
    except Exception as e:
        print(f"✗ 处理失败: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    print("=" * 60)
    print("Markdown到Word格式清理转换器")
    print("=" * 60)
    
    # 处理补充章节
    success = process_supplement_chapter()
    
    if success:
        print("\n" + "=" * 60)
        print("转换成功完成！")
        print("生成的文件已移除所有Markdown格式标记：")
        print("• **粗体** → 粗体")
        print("• *斜体* → 斜体")
        print("• `代码` → 代码")
        print("• [链接](url) → 链接")
        print("• # 标题 → 标题")
        print("=" * 60)
    else:
        print("\n转换失败，请检查错误信息。")
