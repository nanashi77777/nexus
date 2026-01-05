#!/usr/bin/env python3
"""
将Markdown文件转换为Word文档
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_markdown_to_docx(md_file_path, docx_file_path):
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 创建Word文档
    doc = Document()
    
    # 添加自定义样式
    styles = doc.styles
    
    # 代码样式
    code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.right_indent = Inches(0.5)
    
    # 表格样式
    table_style = styles.add_style('Table', WD_STYLE_TYPE.PARAGRAPH)
    table_style.font.name = 'Calibri'
    table_style.font.size = Pt(9)
    
    # 处理状态变量
    in_code_block = False
    in_table = False
    code_block_content = []
    table_rows = []
    
    # 处理每一行
    for line in lines:
        line = line.rstrip('\n')
        
        # 处理代码块开始/结束
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                code_para = doc.add_paragraph(style='Code')
                code_para.add_run('\n'.join(code_block_content))
                in_code_block = False
                code_block_content = []
            else:
                # 开始代码块
                in_code_block = True
            continue
        
        # 如果在代码块中，收集内容
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # 处理表格行
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            continue
        elif in_table:
            # 表格结束，处理表格
            _process_table(doc, table_rows)
            in_table = False
            table_rows = []
        
        # 处理标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=min(level, 6))
            continue
        
        # 处理列表项
        if re.match(r'^[\s]*[-*+]\s+', line):
            # 无序列表
            text = re.sub(r'^[\s]*[-*+]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
            continue
        
        if re.match(r'^[\s]*\d+\.\s+', line):
            # 有序列表
            text = re.sub(r'^[\s]*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
            continue
        
        # 处理空行
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # 普通段落
        p = doc.add_paragraph(line)
    
    # 处理最后可能存在的表格
    if in_table:
        _process_table(doc, table_rows)
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"已成功将 {md_file_path} 转换为 {docx_file_path}")

def _process_table(doc, table_rows):
    """处理Markdown表格"""
    if not table_rows:
        return
    
    # 解析表格内容
    rows = []
    for row in table_rows:
        # 移除前后的|，然后分割
        cells = [cell.strip() for cell in row.strip('|').split('|')]
        rows.append(cells)
    
    if len(rows) < 2:
        return
    
    # 创建Word表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    # 填充表格内容
    for i, row_cells in enumerate(rows):
        for j, cell_content in enumerate(row_cells):
            table.cell(i, j).text = cell_content
    
    # 添加空行分隔
    doc.add_paragraph()

if __name__ == '__main__':
    # 转换文件
    md_file = '技术方案_补充章节.md'
    docx_file = '技术方案_补充章节.docx'
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"转换过程中出现错误: {e}")
        import traceback
        traceback.print_exc()
